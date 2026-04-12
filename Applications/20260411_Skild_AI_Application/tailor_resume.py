"""
Tailor Alexander-Nettekoven-Resume-Skild-AI-Research-Engineer.docx
for Skild AI Research Engineer role.

Changes applied (highlighted in yellow):
  1. Summary [6]: [job title] -> Research Engineer, [company] -> Skild AI
  2. Summary [6]: 'full stack' -> 'scalable'
  3. Summary [6]: 'real-world environment' -> 'real-world environments'
  4. Skills [53] Robot Learning: 'Machine Learning' -> 'Deep Learning'
  5. Skills [55] Robotics: 'Computer Vision' -> 'Perception and Computer Vision'
  6. Skills [57] Programming: add 'large-scale training pipelines' after 'PyTorch'
  7. Skills [58] Team: append '; Cross-Functional Research Collaboration'
"""

import copy
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from lxml import etree

DOCX_IN = "Alexander-Nettekoven-Resume-Skild-AI-Research-Engineer.docx"
DOCX_OUT = "Alexander-Nettekoven-Resume-Skild-AI-Research-Engineer.docx"


def get_full_text(para):
    return "".join(r.text for r in para.runs)


def replace_in_single_run_para(para, old: str, new: str, highlight_words: list[str]):
    """
    For a paragraph whose text lives in runs[0], replace old->new and
    highlight the given highlight_words with yellow.
    The paragraph may have a trailing empty run[1]; we keep it.
    """
    full = get_full_text(para)
    assert old in full, f"Could not find {repr(old)} in {repr(full[:120])}"
    new_full = full.replace(old, new, 1)

    # Reference run for font properties
    ref = para.runs[0]
    ref_size = ref.font.size
    ref_bold = ref.bold
    ref_name = ref.font.name

    # Remove all runs (keep XML elements we'll repopulate)
    for run in para.runs:
        run.text = ""

    # Split new_full around every highlight word in order of appearance
    # Build list of segments: (text, should_highlight)
    segments: list[tuple[str, bool]] = []
    remaining = new_full
    for hw in sorted(highlight_words, key=lambda w: new_full.index(w) if w in new_full else len(new_full)):
        if hw not in remaining:
            continue
        idx = remaining.index(hw)
        if idx > 0:
            segments.append((remaining[:idx], False))
        segments.append((hw, True))
        remaining = remaining[idx + len(hw):]
    if remaining:
        segments.append((remaining, False))

    # Write segments into runs
    # First segment goes into runs[0] to preserve existing XML formatting
    if not segments:
        para.runs[0].text = new_full
        return

    # Assign first segment to run[0]
    first_text, first_hl = segments[0]
    para.runs[0].text = first_text
    if first_hl:
        para.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Remaining segments become new runs appended to paragraph
    for text, hl in segments[1:]:
        new_run = para.add_run(text)
        new_run.font.size = ref_size
        new_run.bold = ref_bold
        new_run.font.name = ref_name
        if hl:
            new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def main():
    doc = Document(DOCX_IN)

    # ---- Edit 1-3: Summary paragraph [6] ----
    p6 = doc.paragraphs[6]
    # Apply all three substitutions in the summary at once
    replace_in_single_run_para(
        p6,
        old="[job title] at [company]. I specialize in robot learning, including VLAs and reinforcement learning, sim-to-real transfer, and full stack robot deployment. 5+ years building novel robotic systems and robot foundation models for real-world environment.",
        new="Research Engineer at Skild AI. I specialize in robot learning, including VLAs and reinforcement learning, sim-to-real transfer, and scalable robot deployment. 5+ years building novel robotic systems and robot foundation models for real-world environments.",
        highlight_words=["Research Engineer", "Skild AI", "scalable", "environments"],
    )

    # ---- Edit 4: Robot Learning skills [53] ----
    p53 = doc.paragraphs[53]
    replace_in_single_run_para(
        p53,
        old="Machine Learning",
        new="Deep Learning",
        highlight_words=["Deep Learning"],
    )

    # ---- Edit 5: Robotics skills [55] ----
    p55 = doc.paragraphs[55]
    replace_in_single_run_para(
        p55,
        old="Computer Vision (SLAM",
        new="Perception and Computer Vision (SLAM",
        highlight_words=["Perception and "],
    )

    # ---- Edit 6: Programming skills [57] ----
    p57 = doc.paragraphs[57]
    replace_in_single_run_para(
        p57,
        old="PyTorch, ROS2",
        new="PyTorch, large-scale training pipelines; ROS2",
        highlight_words=["large-scale training pipelines; "],
    )

    # ---- Edit 7: Team skills [58] ----
    p58 = doc.paragraphs[58]
    full58 = get_full_text(p58)
    replace_in_single_run_para(
        p58,
        old=full58,
        new=full58 + "; Cross-Functional Research Collaboration",
        highlight_words=["; Cross-Functional Research Collaboration"],
    )

    doc.save(DOCX_OUT)
    print("Saved:", DOCX_OUT)

    # Verify
    doc2 = Document(DOCX_OUT)
    print("\nVerification — changed paragraphs:")
    for idx in [6, 53, 55, 57, 58]:
        print(f"  [{idx}] {doc2.paragraphs[idx].text[:110]}")


if __name__ == "__main__":
    main()
